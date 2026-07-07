# -*- coding: utf-8 -*-
r"""
Stage7B_Literature_Filter_and_Novelty_Map.py

Purpose
-------
Refine the large Stage 7A crawler output into a high-priority literature-review set.

Why this stage is needed
------------------------
Stage 7A is intentionally broad and may collect thousands of papers. This script:
1. removes obvious noise,
2. categorizes papers by anatomical site and methodological relevance,
3. prioritizes papers closest to the current HandDx-200 anemia imaging problem,
4. creates a compact Word file for novelty checking,
5. creates CSV tables for manual screening and literature-review planning.

Input
-----
D:\47\472\New-Papers\Anemia_Paper\Outputs\Papers\tables\collected_literature_records.csv

Outputs
-------
D:\47\472\New-Papers\Anemia_Paper\Outputs\Papers\Stage7B_Filtered_Novelty_Map
    ├── tables
    │   ├── Stage7B_all_records_scored.csv
    │   ├── Stage7B_priority_records.csv
    │   ├── Stage7B_direct_hand_related_records.csv
    │   ├── Stage7B_reviews_and_datasets.csv
    │   ├── Stage7B_excluded_low_relevance_records.csv
    │   └── Stage7B_category_summary.csv
    └── reports
        ├── Stage7B_Filtered_Novelty_Map_Report.md
        └── Stage7B_Filtered_Novelty_Map.docx

Recommended install
-------------------
pip install pandas python-docx openpyxl

Run
---
python Stage7B_Literature_Filter_and_Novelty_Map.py
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# PATHS
# ============================================================

INPUT_CSV = Path(
    r"D:\47\472\New-Papers\Anemia_Paper\Outputs\Papers\tables\collected_literature_records.csv"
)

OUTPUT_DIR = Path(
    r"D:\47\472\New-Papers\Anemia_Paper\Outputs\Papers\Stage7B_Filtered_Novelty_Map"
)

TABLES_DIR = OUTPUT_DIR / "tables"
REPORTS_DIR = OUTPUT_DIR / "reports"

for d in [OUTPUT_DIR, TABLES_DIR, REPORTS_DIR]:
    d.mkdir(parents=True, exist_ok=True)


# ============================================================
# CONFIG
# ============================================================

TOP_N_FOR_WORD = 150

# Minimum scores for priority sets.
MIN_PRIORITY_SCORE = 20
MIN_DIRECT_HAND_SCORE = 18

# Records without useful abstracts are usually weak for literature synthesis,
# except when they are clearly known/relevant by title.
REQUIRE_ABSTRACT_FOR_GENERAL_PRIORITY = True


# ============================================================
# TEXT HELPERS
# ============================================================

def norm(x) -> str:
    if pd.isna(x):
        return ""
    x = str(x)
    x = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\u007F-\u009F]", " ", x)
    x = re.sub(r"\s+", " ", x).strip()
    return x


def contains_any(text: str, terms: List[str]) -> bool:
    t = text.lower()
    return any(term.lower() in t for term in terms)


def count_any(text: str, terms: List[str]) -> int:
    t = text.lower()
    return sum(1 for term in terms if term.lower() in t)


def safe_int(x, default=0) -> int:
    try:
        if pd.isna(x):
            return default
        return int(float(x))
    except Exception:
        return default


def safe_float(x, default=0.0) -> float:
    try:
        if pd.isna(x):
            return default
        return float(x)
    except Exception:
        return default


def truncate(text: str, max_chars: int = 2200) -> str:
    text = norm(text)
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + " [...]"


# ============================================================
# KEYWORDS
# ============================================================

ANEMIA_TERMS = [
    "anemia", "anaemia", "hemoglobin", "haemoglobin", "hb", "pallor",
    "iron deficiency", "iron-deficiency"
]

IMAGE_TERMS = [
    "image", "images", "imaging", "photograph", "photo", "camera", "smartphone",
    "computer vision", "deep learning", "machine learning", "cnn", "convolutional",
    "segmentation", "classification", "regression", "estimation"
]

HAND_TERMS = [
    "hand", "hands", "palm", "palmar", "finger", "fingertip", "fingernail",
    "nail", "nailbed", "nail bed", "palmprint"
]

THERMAL_TERMS = [
    "thermal", "infrared", "temperature", "thermography", "perfusion"
]

CONJUNCTIVA_TERMS = [
    "conjunctiva", "conjunctival", "palpebral", "eyelid", "eye"
]

OTHER_BODY_TERMS = [
    "lip", "mucosa", "tongue", "face", "skin", "gingiva"
]

MULTIMODAL_TERMS = [
    "multimodal", "multi-modal", "fusion", "body parts", "multiple body", "multi body",
    "palm and fingernail", "conjunctiva and palm", "palm, and fingernail"
]

DATASET_TERMS = [
    "dataset", "data set", "benchmark", "publicly available", "database"
]

REVIEW_TERMS = [
    "review", "scoping review", "systematic review", "survey"
]

PROTOCOL_TERMS = [
    "protocol", "study protocol"
]

EXCLUDE_STRONG_TERMS = [
    # image-bank examples and microscope/blood-smear papers are not the current problem.
    "ash image bank",
    "blood smear",
    "plasmodium",
    "malaria",
    "microscopy",
    "histopathology",
    "bone marrow",
    "refractory anemia",
    "spur cell anemia",
    "megaloblastic anemia",
    "aplastic anemia",
    "pernicious anemia",
    "sickle cell anemia"
]

CLINICAL_NON_IMAGE_TERMS = [
    "survey analysis",
    "demographic and health survey",
    "national survey",
    "questionnaire",
    "risk factors",
]


# ============================================================
# CLASSIFICATION AND SCORING
# ============================================================

def classify_category(row) -> str:
    text = f"{norm(row.get('title'))} {norm(row.get('abstract'))}".lower()

    if contains_any(text, REVIEW_TERMS):
        return "Review / Survey"

    if contains_any(text, DATASET_TERMS) and contains_any(text, HAND_TERMS):
        return "Dataset / Benchmark"

    if contains_any(text, HAND_TERMS) and contains_any(text, THERMAL_TERMS):
        return "Hand/Palm/Nail + Thermal"

    if contains_any(text, HAND_TERMS) and contains_any(text, MULTIMODAL_TERMS):
        return "Multi-site including Hand/Palm/Nail"

    if contains_any(text, HAND_TERMS):
        return "Direct Hand/Palm/Finger/Nail"

    if contains_any(text, CONJUNCTIVA_TERMS):
        return "Conjunctiva/Eyelid"

    if contains_any(text, OTHER_BODY_TERMS):
        return "Other Visible Body Site"

    if contains_any(text, IMAGE_TERMS):
        return "General Image-based Anemia/Hb"

    return "Low relevance / unclear"


def score_record(row) -> int:
    title = norm(row.get("title"))
    abstract = norm(row.get("abstract"))
    journal = norm(row.get("journal"))
    source_query = norm(row.get("source_query"))
    text = f"{title} {abstract} {journal} {source_query}".lower()

    score = 0

    # Core problem alignment.
    score += 5 * count_any(text, ["anemia", "anaemia"])
    score += 5 * count_any(text, ["hemoglobin", "haemoglobin"])
    score += 3 * count_any(text, ["pallor", "iron deficiency", "iron-deficiency"])

    # Imaging and AI alignment.
    score += 3 * count_any(text, ["image", "imaging", "photograph", "photo", "camera", "smartphone"])
    score += 3 * count_any(text, ["computer vision", "deep learning", "machine learning", "cnn", "convolutional", "classification", "regression", "estimation"])

    # Direct anatomical alignment.
    score += 6 * count_any(text, ["hand", "palm", "palmar"])
    score += 6 * count_any(text, ["finger", "fingertip", "fingernail", "nailbed", "nail bed", "nail"])
    score += 4 * count_any(text, ["conjunctiva", "conjunctival", "palpebral", "eyelid"])
    score += 2 * count_any(text, ["lip", "mucosa", "tongue", "skin", "gingiva"])

    # Current novelty alignment.
    score += 6 * count_any(text, ["thermal", "infrared", "thermography"])
    score += 5 * count_any(text, ["multimodal", "multi-modal", "fusion", "body parts", "multiple body"])

    # Useful literature-review records.
    score += 6 if contains_any(text, REVIEW_TERMS) else 0
    score += 5 if contains_any(text, DATASET_TERMS) else 0

    # Evidence quality proxies.
    citation_count = safe_int(row.get("citation_count"), 0)
    if citation_count >= 100:
        score += 8
    elif citation_count >= 50:
        score += 6
    elif citation_count >= 20:
        score += 4
    elif citation_count >= 5:
        score += 2

    year = safe_int(row.get("year"), 0)
    if year >= 2024:
        score += 4
    elif year >= 2020:
        score += 2

    # Penalties.
    if contains_any(text, EXCLUDE_STRONG_TERMS):
        score -= 30

    if contains_any(text, CLINICAL_NON_IMAGE_TERMS) and not contains_any(text, ["image", "photo", "camera", "vision", "smartphone"]):
        score -= 12

    if not abstract and not contains_any(title.lower(), ["image", "imaging", "smartphone", "palm", "finger", "nail", "conjunctiva", "thermal"]):
        score -= 15

    return int(score)


def assign_decision(row) -> str:
    score = safe_int(row.get("priority_score"))
    category = norm(row.get("category"))
    abstract = norm(row.get("abstract"))
    text = f"{norm(row.get('title'))} {abstract}".lower()

    if contains_any(text, EXCLUDE_STRONG_TERMS):
        return "Exclude - not current problem"

    if category in [
        "Direct Hand/Palm/Finger/Nail",
        "Hand/Palm/Nail + Thermal",
        "Multi-site including Hand/Palm/Nail",
        "Dataset / Benchmark",
    ] and score >= MIN_DIRECT_HAND_SCORE:
        return "Priority A - Directly related"

    if category in ["Conjunctiva/Eyelid", "Other Visible Body Site", "General Image-based Anemia/Hb"] and score >= MIN_PRIORITY_SCORE:
        if REQUIRE_ABSTRACT_FOR_GENERAL_PRIORITY and not abstract:
            return "Priority C - metadata only"
        return "Priority B - Related background"

    if category == "Review / Survey" and score >= MIN_PRIORITY_SCORE:
        return "Priority A - Review foundation"

    if "protocol" in text:
        return "Priority C - protocol / future evidence"

    if score >= MIN_PRIORITY_SCORE and abstract:
        return "Priority C - check manually"

    return "Exclude / low priority"


# ============================================================
# WORD REPORT
# ============================================================

def add_para(doc, text, size=10, bold=False):
    text = truncate(text, 4000)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_field(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(truncate(value, 1500) if value else "Not available")
    r2.font.size = Pt(9)


def create_word_report(priority_df: pd.DataFrame, summary_df: pd.DataFrame, output_path: Path):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_heading("Stage 7B Filtered Novelty Map for Image-Based Anemia Detection", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "This document condenses the broad crawler output into a prioritized evidence base for novelty assessment, "
        "introduction development, and literature-review planning. Records are ranked according to their relevance to "
        "non-invasive anemia or hemoglobin assessment using hand, palm, finger, fingernail, conjunctiva, RGB, thermal, "
        "smartphone, computer-vision, and deep-learning methods."
    )

    doc.add_heading("Category Summary", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Records"
    hdr[2].text = "Priority A"
    hdr[3].text = "Priority B/C"

    for _, row in summary_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["category"])
        cells[1].text = str(row["records"])
        cells[2].text = str(row["priority_a"])
        cells[3].text = str(row["priority_b_c"])

    doc.add_heading("High-Priority Papers", level=1)

    for idx, (_, r) in enumerate(priority_df.head(TOP_N_FOR_WORD).iterrows(), start=1):
        doc.add_heading(f"{idx}. {norm(r.get('title'))}", level=2)

        add_field(doc, "Decision", norm(r.get("screening_decision")))
        add_field(doc, "Category", norm(r.get("category")))
        add_field(doc, "Priority score", str(r.get("priority_score")))
        add_field(doc, "Year", str(r.get("year")))
        add_field(doc, "Journal", norm(r.get("journal")))
        add_field(doc, "Authors", norm(r.get("authors")))
        add_field(doc, "DOI", norm(r.get("doi")))
        add_field(doc, "URL", norm(r.get("url")))
        add_field(doc, "PDF path", norm(r.get("pdf_path")))
        add_field(doc, "Citation count", str(r.get("citation_count")))

        doc.add_heading("Abstract", level=3)
        add_para(doc, norm(r.get("abstract")) or "Not available.", size=9)

        doc.add_heading("Conclusion / Discussion Excerpt", level=3)
        add_para(doc, norm(r.get("conclusion")) or "Not available or not automatically detected.", size=9)

        doc.add_heading("Citation", level=3)
        add_para(doc, norm(r.get("citation")), size=9)

    doc.save(output_path)


# ============================================================
# MAIN
# ============================================================

def main():
    if not INPUT_CSV.exists():
        raise FileNotFoundError(f"Missing input CSV: {INPUT_CSV}")

    df = pd.read_csv(INPUT_CSV)

    for col in ["title", "abstract", "authors", "journal", "doi", "url", "pdf_url", "source", "source_query", "pdf_path", "conclusion", "citation"]:
        if col not in df.columns:
            df[col] = ""

    df["title"] = df["title"].map(norm)
    df["abstract"] = df["abstract"].map(norm)
    df["journal"] = df["journal"].map(norm)
    df["citation"] = df["citation"].map(norm)
    df["conclusion"] = df["conclusion"].map(norm)

    df["category"] = df.apply(classify_category, axis=1)
    df["priority_score"] = df.apply(score_record, axis=1)
    df["screening_decision"] = df.apply(assign_decision, axis=1)

    # Sort by decision and score.
    decision_rank = {
        "Priority A - Directly related": 1,
        "Priority A - Review foundation": 2,
        "Priority B - Related background": 3,
        "Priority C - protocol / future evidence": 4,
        "Priority C - check manually": 5,
        "Priority C - metadata only": 6,
        "Exclude - not current problem": 9,
        "Exclude / low priority": 10,
    }
    df["decision_rank"] = df["screening_decision"].map(decision_rank).fillna(8).astype(int)

    df = df.sort_values(
        by=["decision_rank", "priority_score", "citation_count", "year"],
        ascending=[True, False, False, False]
    )

    priority_df = df[df["screening_decision"].str.startswith("Priority")].copy()
    direct_df = df[df["screening_decision"].eq("Priority A - Directly related")].copy()
    review_dataset_df = df[
        df["category"].isin(["Review / Survey", "Dataset / Benchmark"])
        & df["screening_decision"].str.startswith("Priority")
    ].copy()
    excluded_df = df[df["screening_decision"].str.startswith("Exclude")].copy()

    summary_rows = []
    for cat, g in df.groupby("category"):
        summary_rows.append({
            "category": cat,
            "records": len(g),
            "priority_a": int(g["screening_decision"].str.startswith("Priority A").sum()),
            "priority_b_c": int(
                g["screening_decision"].str.startswith("Priority B").sum()
                + g["screening_decision"].str.startswith("Priority C").sum()
            ),
            "excluded": int(g["screening_decision"].str.startswith("Exclude").sum()),
            "mean_priority_score": round(g["priority_score"].mean(), 2),
        })

    summary_df = pd.DataFrame(summary_rows).sort_values(
        by=["priority_a", "priority_b_c", "records"],
        ascending=False
    )

    all_path = TABLES_DIR / "Stage7B_all_records_scored.csv"
    priority_path = TABLES_DIR / "Stage7B_priority_records.csv"
    direct_path = TABLES_DIR / "Stage7B_direct_hand_related_records.csv"
    review_dataset_path = TABLES_DIR / "Stage7B_reviews_and_datasets.csv"
    excluded_path = TABLES_DIR / "Stage7B_excluded_low_relevance_records.csv"
    summary_path = TABLES_DIR / "Stage7B_category_summary.csv"

    df.to_csv(all_path, index=False, encoding="utf-8-sig")
    priority_df.to_csv(priority_path, index=False, encoding="utf-8-sig")
    direct_df.to_csv(direct_path, index=False, encoding="utf-8-sig")
    review_dataset_df.to_csv(review_dataset_path, index=False, encoding="utf-8-sig")
    excluded_df.to_csv(excluded_path, index=False, encoding="utf-8-sig")
    summary_df.to_csv(summary_path, index=False, encoding="utf-8-sig")

    word_path = REPORTS_DIR / "Stage7B_Filtered_Novelty_Map.docx"
    create_word_report(priority_df, summary_df, word_path)

    md = f"""# Stage 7B Filtered Novelty Map Report

## Input
`{INPUT_CSV}`

## Records
- Total Stage 7A records: {len(df)}
- Priority records retained: {len(priority_df)}
- Direct hand/palm/finger/nail records: {len(direct_df)}
- Review/dataset foundation records: {len(review_dataset_df)}
- Excluded or low-priority records: {len(excluded_df)}

## Output files
- All scored records: `{all_path}`
- Priority records: `{priority_path}`
- Direct hand-related records: `{direct_path}`
- Reviews and datasets: `{review_dataset_path}`
- Excluded records: `{excluded_path}`
- Category summary: `{summary_path}`
- Word novelty map: `{word_path}`

## Interpretation
Priority A records should be read first because they directly address non-invasive anemia or hemoglobin assessment using hand, palm, finger, fingernail, nailbed, thermal, RGB, or multimodal image evidence. Priority B records are useful for background because they address image-based anemia detection from conjunctiva, lip mucosa, or other visible body sites. Priority C records may be useful for context, protocols, or manual checking.

The strongest novelty direction for the current HandDx-200 work should be assessed against the Priority A group. Particular attention should be given to whether prior studies used:
1. true participant-level multi-view hand modeling,
2. combined RGB and thermal hand imaging,
3. view-level attention or anatomical-view fusion,
4. explainable model behavior,
5. external validation or robustness analysis.
"""
    (REPORTS_DIR / "Stage7B_Filtered_Novelty_Map_Report.md").write_text(md, encoding="utf-8")

    print("=" * 90)
    print("STAGE 7B FILTERED NOVELTY MAP COMPLETED")
    print("=" * 90)
    print(f"Total records scored: {len(df)}")
    print(f"Priority records retained: {len(priority_df)}")
    print(f"Direct hand-related records: {len(direct_df)}")
    print(f"Review/dataset foundation records: {len(review_dataset_df)}")
    print(f"Excluded / low-priority records: {len(excluded_df)}")
    print(f"Word novelty map: {word_path}")
    print(f"Priority CSV: {priority_path}")
    print("=" * 90)


if __name__ == "__main__":
    main()
