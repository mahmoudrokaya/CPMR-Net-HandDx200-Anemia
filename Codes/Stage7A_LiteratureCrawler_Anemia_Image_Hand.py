# -*- coding: utf-8 -*-
"""
Stage7A_LiteratureCrawler_Anemia_Image_Hand.py

Purpose
-------
Collect reputable scholarly papers related to non-invasive anemia detection/classification
using images in general and hand/finger/palm/conjunctiva/nail images specifically.

The script searches multiple scholarly sources, downloads legally available open-access PDFs
when possible, extracts title, abstract, conclusion/discussion text, and creates one Word file
that can be used as a novelty-checking base for the introduction and literature review.

Sources used
------------
1. Europe PMC
2. OpenAlex
3. Semantic Scholar
4. Crossref

Important notes
---------------
- The script does NOT bypass paywalls.
- PDF download is attempted only from open URLs returned by reputable metadata APIs.
- Some papers may have metadata only, without a downloadable PDF.
- Conclusion extraction is heuristic because publishers format PDFs differently.

Recommended install
-------------------
pip install requests pandas python-docx pymupdf tqdm beautifulsoup4 lxml

Run
---
python Stage7A_LiteratureCrawler_Anemia_Image_Hand.py
"""

from __future__ import annotations

import json
import re
import time
import hashlib
import textwrap
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Optional, Iterable, Tuple
from urllib.parse import quote_plus, urlparse

import requests
import pandas as pd
from tqdm import tqdm
from bs4 import BeautifulSoup

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# PATHS
# ============================================================

OUTPUT_ROOT = Path(r"D:\47\472\New-Papers\Anemia_Paper\Outputs\Papers")
PDF_DIR = OUTPUT_ROOT / "pdfs"
TABLES_DIR = OUTPUT_ROOT / "tables"
REPORTS_DIR = OUTPUT_ROOT / "reports"
RAW_DIR = OUTPUT_ROOT / "raw_api_results"

for d in [OUTPUT_ROOT, PDF_DIR, TABLES_DIR, REPORTS_DIR, RAW_DIR]:
    d.mkdir(parents=True, exist_ok=True)


# ============================================================
# CONFIG
# ============================================================

USER_EMAIL = "baselmah@yahoo.com"  # used only in polite API headers / Unpaywall-compatible future extension
MAX_RESULTS_PER_QUERY_PER_SOURCE = 50
REQUEST_TIMEOUT = 30
SLEEP_BETWEEN_REQUESTS = 0.5
DOWNLOAD_PDFS = True
MAX_PDF_MB = 80

# Strong but flexible query set.
# It intentionally covers anemia, hemoglobin, pallor, image-based diagnosis,
# hands/fingers/palms/nails/conjunctiva, RGB, thermal, smartphone, deep learning.
QUERIES = [
    '"anemia" AND image AND classification',
    '"anaemia" AND image AND classification',
    '"anemia detection" AND image',
    '"anaemia detection" AND image',
    '"non-invasive" AND anemia AND imaging',
    '"noninvasive" AND anemia AND imaging',
    '"hemoglobin" AND image AND estimation',
    '"hemoglobin" AND smartphone AND image',
    '"hemoglobin concentration" AND image',
    '"pallor" AND image AND anemia',
    '"palmar pallor" AND anemia',
    '"hand image" AND anemia',
    '"hand images" AND hemoglobin',
    '"palm image" AND anemia',
    '"finger image" AND anemia',
    '"fingernail" AND anemia AND image',
    '"nailbed" AND anemia AND image',
    '"conjunctiva" AND anemia AND image',
    '"conjunctival pallor" AND anemia AND image',
    '"RGB" AND anemia AND image',
    '"thermal imaging" AND anemia',
    '"infrared" AND anemia AND imaging',
    '"deep learning" AND anemia AND image',
    '"machine learning" AND anemia AND image',
    '"computer vision" AND anemia',
    '"multimodal" AND anemia AND imaging',
    '"hand" AND "thermal" AND "anemia"',
    '"hand" AND "RGB" AND "anemia"',
    '"HandDx" OR "HandDx-200"',
]

# Terms used after collection to score relevance to the current work.
HIGH_RELEVANCE_TERMS = [
    "anemia", "anaemia", "hemoglobin", "haemoglobin", "pallor",
    "hand", "palm", "palmar", "finger", "fingernail", "nailbed",
    "conjunctiva", "conjunctival", "rgb", "thermal", "infrared",
    "image", "imaging", "smartphone", "computer vision", "deep learning",
    "machine learning", "non-invasive", "noninvasive", "classification",
]


# ============================================================
# HTTP HELPERS
# ============================================================

HEADERS = {
    "User-Agent": f"AnemiaLiteratureCrawler/1.0 (mailto:{USER_EMAIL})",
    "Accept": "application/json,text/html,application/pdf;q=0.9,*/*;q=0.8",
}


def safe_get(url: str, params: Optional[dict] = None, headers: Optional[dict] = None,
             stream: bool = False) -> Optional[requests.Response]:
    try:
        time.sleep(SLEEP_BETWEEN_REQUESTS)
        r = requests.get(
            url,
            params=params,
            headers=headers or HEADERS,
            timeout=REQUEST_TIMEOUT,
            stream=stream,
            allow_redirects=True,
        )
        if r.status_code in {200, 201}:
            return r
        return None
    except Exception:
        return None


def normalize_space(x: Optional[str]) -> str:
    if not x:
        return ""
    return re.sub(r"\s+", " ", str(x)).strip()


def clean_html(x: Optional[str]) -> str:
    if not x:
        return ""
    soup = BeautifulSoup(str(x), "lxml")
    return normalize_space(soup.get_text(" "))


def doi_to_url(doi: str) -> str:
    doi = normalize_space(doi)
    if not doi:
        return ""
    doi = re.sub(r"^https?://(dx\.)?doi\.org/", "", doi, flags=re.I)
    return f"https://doi.org/{doi}"


def stable_id(*parts: str) -> str:
    text = "||".join([normalize_space(p).lower() for p in parts if p])
    return hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()[:16]


# ============================================================
# DATA STRUCTURE
# ============================================================

@dataclass
class PaperRecord:
    paper_id: str
    title: str = ""
    abstract: str = ""
    authors: str = ""
    year: str = ""
    journal: str = ""
    doi: str = ""
    url: str = ""
    pdf_url: str = ""
    source: str = ""
    source_query: str = ""
    citation_count: Optional[int] = None
    open_access: Optional[bool] = None
    relevance_score: int = 0
    pdf_path: str = ""
    conclusion: str = ""
    citation: str = ""
    notes: str = ""


# ============================================================
# SOURCE SEARCHERS
# ============================================================

def search_europe_pmc(query: str, page_size: int = 50) -> List[PaperRecord]:
    url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
    params = {
        "query": query,
        "format": "json",
        "pageSize": min(page_size, 100),
        "resultType": "core",
    }
    r = safe_get(url, params=params)
    if not r:
        return []

    raw_path = RAW_DIR / f"europepmc_{stable_id(query)}.json"
    raw_path.write_text(r.text, encoding="utf-8")

    data = r.json()
    out = []
    for item in data.get("resultList", {}).get("result", []):
        title = clean_html(item.get("title"))
        abstract = clean_html(item.get("abstractText"))
        doi = normalize_space(item.get("doi"))
        year = normalize_space(item.get("pubYear"))
        journal = normalize_space(item.get("journalTitle"))
        authors = normalize_space(item.get("authorString"))
        cited = item.get("citedByCount")
        pmcid = normalize_space(item.get("pmcid"))
        pmid = normalize_space(item.get("pmid"))

        pdf_url = ""
        if pmcid:
            pdf_url = f"https://europepmc.org/articles/{pmcid}?pdf=render"

        url_best = doi_to_url(doi) if doi else (f"https://europepmc.org/article/MED/{pmid}" if pmid else "")
        pid = stable_id(doi, title, year)
        out.append(PaperRecord(
            paper_id=pid,
            title=title,
            abstract=abstract,
            authors=authors,
            year=year,
            journal=journal,
            doi=doi,
            url=url_best,
            pdf_url=pdf_url,
            source="Europe PMC",
            source_query=query,
            citation_count=int(cited) if str(cited).isdigit() else None,
            open_access=bool(pmcid),
        ))
    return out


def search_openalex(query: str, per_page: int = 50) -> List[PaperRecord]:
    url = "https://api.openalex.org/works"
    params = {
        "search": query.replace('"', ""),
        "per-page": min(per_page, 200),
        "mailto": USER_EMAIL,
    }
    r = safe_get(url, params=params)
    if not r:
        return []

    raw_path = RAW_DIR / f"openalex_{stable_id(query)}.json"
    raw_path.write_text(r.text, encoding="utf-8")

    data = r.json()
    out = []
    for item in data.get("results", []):
        title = clean_html(item.get("title") or item.get("display_name"))
        doi = normalize_space(item.get("doi") or "")
        doi = re.sub(r"^https?://doi\.org/", "", doi, flags=re.I)
        year = str(item.get("publication_year") or "")
        journal = normalize_space(
            ((item.get("primary_location") or {}).get("source") or {}).get("display_name")
        )
        authors = "; ".join([
            normalize_space((a.get("author") or {}).get("display_name"))
            for a in item.get("authorships", [])[:12]
            if (a.get("author") or {}).get("display_name")
        ])
        if len(item.get("authorships", [])) > 12:
            authors += "; et al."

        abstract = reconstruct_openalex_abstract(item.get("abstract_inverted_index"))
        oa = item.get("open_access") or {}
        best_oa = (item.get("best_oa_location") or {})
        pdf_url = normalize_space(best_oa.get("pdf_url"))
        landing = normalize_space(best_oa.get("landing_page_url")) or doi_to_url(doi) or normalize_space(item.get("id"))
        cited = item.get("cited_by_count")

        pid = stable_id(doi, title, year)
        out.append(PaperRecord(
            paper_id=pid,
            title=title,
            abstract=abstract,
            authors=authors,
            year=year,
            journal=journal,
            doi=doi,
            url=landing,
            pdf_url=pdf_url,
            source="OpenAlex",
            source_query=query,
            citation_count=int(cited) if isinstance(cited, int) else None,
            open_access=bool(oa.get("is_oa")),
        ))
    return out


def reconstruct_openalex_abstract(inv: Optional[dict]) -> str:
    if not inv:
        return ""
    positions = []
    for word, idxs in inv.items():
        for idx in idxs:
            positions.append((idx, word))
    positions.sort()
    return normalize_space(" ".join(w for _, w in positions))


def search_semantic_scholar(query: str, limit: int = 50) -> List[PaperRecord]:
    url = "https://api.semanticscholar.org/graph/v1/paper/search"
    fields = ",".join([
        "title", "abstract", "authors", "year", "venue", "externalIds",
        "url", "openAccessPdf", "citationCount"
    ])
    params = {
        "query": query.replace('"', ""),
        "limit": min(limit, 100),
        "fields": fields,
    }
    r = safe_get(url, params=params)
    if not r:
        return []

    raw_path = RAW_DIR / f"semanticscholar_{stable_id(query)}.json"
    raw_path.write_text(r.text, encoding="utf-8")

    data = r.json()
    out = []
    for item in data.get("data", []):
        title = clean_html(item.get("title"))
        abstract = clean_html(item.get("abstract"))
        year = str(item.get("year") or "")
        journal = normalize_space(item.get("venue"))
        authors = "; ".join([normalize_space(a.get("name")) for a in item.get("authors", [])[:12]])
        if len(item.get("authors", [])) > 12:
            authors += "; et al."
        ext = item.get("externalIds") or {}
        doi = normalize_space(ext.get("DOI"))
        oa = item.get("openAccessPdf") or {}
        pdf_url = normalize_space(oa.get("url"))
        cited = item.get("citationCount")
        url_best = normalize_space(item.get("url")) or doi_to_url(doi)

        pid = stable_id(doi, title, year)
        out.append(PaperRecord(
            paper_id=pid,
            title=title,
            abstract=abstract,
            authors=authors,
            year=year,
            journal=journal,
            doi=doi,
            url=url_best,
            pdf_url=pdf_url,
            source="Semantic Scholar",
            source_query=query,
            citation_count=int(cited) if isinstance(cited, int) else None,
            open_access=bool(pdf_url),
        ))
    return out


def search_crossref(query: str, rows: int = 50) -> List[PaperRecord]:
    url = "https://api.crossref.org/works"
    params = {
        "query": query.replace('"', ""),
        "rows": min(rows, 100),
        "mailto": USER_EMAIL,
        "filter": "type:journal-article",
    }
    r = safe_get(url, params=params)
    if not r:
        return []

    raw_path = RAW_DIR / f"crossref_{stable_id(query)}.json"
    raw_path.write_text(r.text, encoding="utf-8")

    data = r.json()
    out = []
    for item in data.get("message", {}).get("items", []):
        title = clean_html(" ".join(item.get("title", [])[:1]))
        abstract = clean_html(item.get("abstract"))
        doi = normalize_space(item.get("DOI"))
        year = ""
        issued = item.get("issued", {}).get("date-parts", [])
        if issued and issued[0]:
            year = str(issued[0][0])
        journal = normalize_space(" ".join(item.get("container-title", [])[:1]))
        authors_list = []
        for a in item.get("author", [])[:12]:
            name = normalize_space(f"{a.get('given', '')} {a.get('family', '')}")
            if name:
                authors_list.append(name)
        authors = "; ".join(authors_list)
        if len(item.get("author", [])) > 12:
            authors += "; et al."

        pdf_url = ""
        for link in item.get("link", []):
            content_type = normalize_space(link.get("content-type")).lower()
            candidate = normalize_space(link.get("URL"))
            if "pdf" in content_type or candidate.lower().endswith(".pdf"):
                pdf_url = candidate
                break

        cited = item.get("is-referenced-by-count")
        pid = stable_id(doi, title, year)
        out.append(PaperRecord(
            paper_id=pid,
            title=title,
            abstract=abstract,
            authors=authors,
            year=year,
            journal=journal,
            doi=doi,
            url=doi_to_url(doi),
            pdf_url=pdf_url,
            source="Crossref",
            source_query=query,
            citation_count=int(cited) if isinstance(cited, int) else None,
            open_access=bool(pdf_url),
        ))
    return out


# ============================================================
# DEDUPLICATION AND SCORING
# ============================================================

def merge_records(records: List[PaperRecord]) -> List[PaperRecord]:
    by_key: Dict[str, PaperRecord] = {}

    def key_for(r: PaperRecord) -> str:
        if r.doi:
            return "doi:" + r.doi.lower().strip()
        return "title:" + re.sub(r"[^a-z0-9]+", "", r.title.lower())[:100]

    for r in records:
        if not r.title:
            continue
        k = key_for(r)
        if k not in by_key:
            by_key[k] = r
        else:
            old = by_key[k]
            old.abstract = old.abstract or r.abstract
            old.authors = old.authors or r.authors
            old.year = old.year or r.year
            old.journal = old.journal or r.journal
            old.doi = old.doi or r.doi
            old.url = old.url or r.url
            old.pdf_url = old.pdf_url or r.pdf_url
            old.open_access = bool(old.open_access or r.open_access)
            old.citation_count = max(
                [x for x in [old.citation_count, r.citation_count] if x is not None],
                default=None,
            )
            old.source = old.source + " | " + r.source if r.source not in old.source else old.source
            if r.source_query not in old.source_query:
                old.source_query += " | " + r.source_query
    return list(by_key.values())


def score_relevance(r: PaperRecord) -> int:
    text = f"{r.title} {r.abstract} {r.source_query}".lower()
    score = 0
    for term in HIGH_RELEVANCE_TERMS:
        if term.lower() in text:
            score += 1

    # Strong boosts for the exact current problem.
    boosts = [
        ("anemia", 4), ("anaemia", 4), ("hemoglobin", 4), ("haemoglobin", 4),
        ("hand", 5), ("palm", 5), ("palmar", 5), ("finger", 4), ("nail", 4),
        ("conjunctiva", 3), ("thermal", 3), ("rgb", 3), ("image", 2),
        ("classification", 2), ("deep learning", 2), ("machine learning", 2),
    ]
    for term, weight in boosts:
        if term in text:
            score += weight

    # Penalize if it is only a general anemia clinical paper without image/computer vision terms.
    if ("anemia" in text or "anaemia" in text) and not any(
        t in text for t in ["image", "imaging", "photograph", "smartphone", "vision", "thermal", "rgb", "conjunctiva", "hand", "palm", "finger", "nail"]
    ):
        score -= 8

    return score


# ============================================================
# PDF DOWNLOAD AND TEXT EXTRACTION
# ============================================================

def safe_filename(text: str, max_len: int = 120) -> str:
    text = normalize_space(text)
    text = re.sub(r"[^\w\s\-\.\(\)]", "_", text, flags=re.UNICODE)
    text = re.sub(r"\s+", "_", text)
    return text[:max_len].strip("_") or "paper"


def is_pdf_response(resp: requests.Response, url: str) -> bool:
    ctype = resp.headers.get("content-type", "").lower()
    return "pdf" in ctype or url.lower().split("?")[0].endswith(".pdf") or resp.content[:4] == b"%PDF"


def download_pdf(r: PaperRecord) -> str:
    if not r.pdf_url:
        return ""
    try:
        resp = safe_get(r.pdf_url, stream=True)
        if not resp:
            return ""
        content_length = resp.headers.get("content-length")
        if content_length and int(content_length) > MAX_PDF_MB * 1024 * 1024:
            return ""
        first_chunk = resp.raw.read(4096, decode_content=True)
        if not first_chunk:
            return ""

        # Re-request normally because we consumed the stream header.
        resp.close()
        resp = safe_get(r.pdf_url)
        if not resp or not is_pdf_response(resp, r.pdf_url):
            return ""

        fname = safe_filename(f"{r.year}_{r.title}") + ".pdf"
        out_path = PDF_DIR / fname
        counter = 1
        while out_path.exists():
            out_path = PDF_DIR / (safe_filename(f"{r.year}_{r.title}") + f"_{counter}.pdf")
            counter += 1

        if len(resp.content) > MAX_PDF_MB * 1024 * 1024:
            return ""

        out_path.write_bytes(resp.content)
        return str(out_path)
    except Exception:
        return ""


def extract_pdf_text(pdf_path: str) -> str:
    if not pdf_path:
        return ""
    try:
        doc = fitz.open(pdf_path)
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return normalize_space("\n".join(pages))
    except Exception:
        return ""


def extract_section(text: str, section_names: Iterable[str], max_chars: int = 7000) -> str:
    if not text:
        return ""
    pattern_names = "|".join([re.escape(s) for s in section_names])
    # Common section heading forms: "4. Conclusion", "Conclusions", "Discussion and conclusion"
    start_re = re.compile(
        rf"(?i)(?:^|\n|\s)(?:\d+\.?\s*)?({pattern_names})\s*(?:\n|$|:)",
        flags=re.I,
    )
    matches = list(start_re.finditer(text))
    if not matches:
        return ""

    start = matches[-1].start()
    next_heading = re.search(
        r"(?i)\n\s*(?:\d+\.?\s*)?(references|acknowledg|funding|conflict|ethics|data availability|supplementary)\s*(?:\n|:)",
        text[start:],
    )
    end = start + next_heading.start() if next_heading else min(len(text), start + max_chars)
    section = normalize_space(text[start:end])
    return section[:max_chars]


def extract_conclusion(text: str) -> str:
    return extract_section(text, [
        "conclusion",
        "conclusions",
        "discussion and conclusion",
        "conclusion and future work",
        "summary and conclusion",
        "summary",
    ])


def extract_abstract_from_pdf(text: str) -> str:
    return extract_section(text, ["abstract"], max_chars=4000)


# ============================================================
# CITATION
# ============================================================

def make_citation(r: PaperRecord) -> str:
    authors = r.authors or "Unknown authors"
    year = r.year or "n.d."
    title = r.title or "Untitled"
    journal = r.journal or "Unknown source"
    doi_part = f" https://doi.org/{r.doi}" if r.doi else (f" {r.url}" if r.url else "")
    return f"{authors} ({year}). {title}. {journal}.{doi_part}"


# ============================================================
# WORD REPORT
# ============================================================

def add_normal_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_field(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(value if value else "Not available")
    r2.font.size = Pt(10)


def create_word_report(records: List[PaperRecord], output_docx: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_heading("Literature Evidence Base for Image-Based Non-Invasive Anemia Classification", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_normal_paragraph(
        doc,
        "This file was automatically generated to support novelty analysis, introduction writing, and literature review development. "
        "It collects scholarly records from reputable metadata sources and includes title, abstract, conclusion/discussion excerpts when available, "
        "PDF availability, and a citation entry for each paper. Open-access PDFs were downloaded only when a legal PDF URL was available from scholarly sources."
    )

    doc.add_heading("Search Strategy", level=1)
    add_normal_paragraph(doc, "Sources: Europe PMC, OpenAlex, Semantic Scholar, and Crossref.")
    add_normal_paragraph(doc, "Main query themes: anemia/anaemia, hemoglobin/haemoglobin, non-invasive imaging, RGB imaging, thermal imaging, hand/palm/finger/fingernail/nailbed/conjunctiva images, smartphone imaging, machine learning, deep learning, and computer vision.")

    doc.add_heading("Summary", level=1)
    add_normal_paragraph(doc, f"Unique records collected: {len(records)}")
    add_normal_paragraph(doc, f"Records with downloaded PDFs: {sum(1 for r in records if r.pdf_path)}")
    add_normal_paragraph(doc, f"Records with abstract text: {sum(1 for r in records if r.abstract)}")
    add_normal_paragraph(doc, f"Records with extracted conclusion/discussion text: {sum(1 for r in records if r.conclusion)}")

    doc.add_heading("Collected Papers", level=1)

    for i, r in enumerate(records, start=1):
        doc.add_heading(f"{i}. {r.title}", level=2)
        add_field(doc, "Authors", r.authors)
        add_field(doc, "Year", r.year)
        add_field(doc, "Journal / Venue", r.journal)
        add_field(doc, "DOI", r.doi)
        add_field(doc, "URL", r.url)
        add_field(doc, "PDF URL", r.pdf_url)
        add_field(doc, "Downloaded PDF", r.pdf_path)
        add_field(doc, "Source", r.source)
        add_field(doc, "Citation count", str(r.citation_count) if r.citation_count is not None else "")
        add_field(doc, "Relevance score", str(r.relevance_score))

        doc.add_heading("Abstract", level=3)
        add_normal_paragraph(doc, r.abstract or "Not available from metadata or PDF extraction.")

        doc.add_heading("Conclusion / Discussion Excerpt", level=3)
        add_normal_paragraph(doc, r.conclusion or "Not available or not automatically detected from the PDF.")

        doc.add_heading("Citation", level=3)
        add_normal_paragraph(doc, r.citation)

        doc.add_paragraph("")

    doc.add_heading("Bibliography", level=1)
    for i, r in enumerate(records, start=1):
        add_normal_paragraph(doc, f"[{i}] {r.citation}")

    doc.save(output_docx)


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 90)
    print("STAGE 7A LITERATURE CRAWLER: IMAGE-BASED ANEMIA DETECTION")
    print("=" * 90)
    print(f"Output folder: {OUTPUT_ROOT}")

    all_records: List[PaperRecord] = []

    source_functions = [
        ("Europe PMC", search_europe_pmc),
        ("OpenAlex", search_openalex),
        ("Semantic Scholar", search_semantic_scholar),
        ("Crossref", search_crossref),
    ]

    for query in QUERIES:
        print(f"\nQuery: {query}")
        for source_name, func in source_functions:
            try:
                print(f"  Searching {source_name}...")
                recs = func(query, MAX_RESULTS_PER_QUERY_PER_SOURCE)
                print(f"    Found: {len(recs)}")
                all_records.extend(recs)
            except Exception as e:
                print(f"    Failed {source_name}: {e}")

    print("\nDeduplicating records...")
    records = merge_records(all_records)

    for r in records:
        r.relevance_score = score_relevance(r)
        r.citation = make_citation(r)

    # Keep relevant records. Threshold intentionally permissive for recall.
    records = [r for r in records if r.relevance_score >= 8]
    records = sorted(
        records,
        key=lambda x: (
            x.relevance_score,
            x.citation_count if x.citation_count is not None else -1,
            int(x.year) if str(x.year).isdigit() else 0,
        ),
        reverse=True,
    )

    print(f"Unique relevant records retained: {len(records)}")

    if DOWNLOAD_PDFS:
        print("\nDownloading open-access PDFs and extracting conclusions...")
        for r in tqdm(records):
            if r.pdf_url:
                r.pdf_path = download_pdf(r)
            pdf_text = extract_pdf_text(r.pdf_path) if r.pdf_path else ""
            if not r.abstract and pdf_text:
                r.abstract = extract_abstract_from_pdf(pdf_text)
            if pdf_text:
                r.conclusion = extract_conclusion(pdf_text)

    # Save metadata tables.
    rows = [asdict(r) for r in records]
    df = pd.DataFrame(rows)
    csv_path = TABLES_DIR / "collected_literature_records.csv"
    xlsx_path = TABLES_DIR / "collected_literature_records.xlsx"
    json_path = TABLES_DIR / "collected_literature_records.json"

    df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    df.to_excel(xlsx_path, index=False)
    json_path.write_text(json.dumps(rows, indent=2, ensure_ascii=False), encoding="utf-8")

    # Create Word report.
    docx_path = OUTPUT_ROOT / "Image_Based_Anemia_Literature_Evidence_Base.docx"
    create_word_report(records, docx_path)

    # Also write run report.
    md_report = f"""# Stage 7A Literature Crawler Report

## Objective
Collect scholarly works that address non-invasive anemia / hemoglobin detection using images in general and hand-related images specifically.

## Sources
- Europe PMC
- OpenAlex
- Semantic Scholar
- Crossref

## Output
- Word evidence file: `{docx_path}`
- Metadata CSV: `{csv_path}`
- Metadata Excel: `{xlsx_path}`
- Raw API responses: `{RAW_DIR}`
- Downloaded PDFs: `{PDF_DIR}`

## Results
- Relevant unique records retained: {len(records)}
- PDFs downloaded: {sum(1 for r in records if r.pdf_path)}
- Records with abstracts: {sum(1 for r in records if r.abstract)}
- Records with extracted conclusion/discussion text: {sum(1 for r in records if r.conclusion)}

## Notes
The crawler does not bypass paywalls. Papers without open-access PDF links are still retained if their metadata is relevant.
Conclusion extraction is heuristic and should be manually checked for high-priority papers.
"""
    (REPORTS_DIR / "Stage7A_LiteratureCrawler_Report.md").write_text(md_report, encoding="utf-8")

    print("=" * 90)
    print("STAGE 7A LITERATURE CRAWLER COMPLETED")
    print("=" * 90)
    print(f"Relevant unique records: {len(records)}")
    print(f"PDFs downloaded: {sum(1 for r in records if r.pdf_path)}")
    print(f"Word file: {docx_path}")
    print(f"CSV file: {csv_path}")
    print(f"Excel file: {xlsx_path}")
    print("=" * 90)


if __name__ == "__main__":
    main()
